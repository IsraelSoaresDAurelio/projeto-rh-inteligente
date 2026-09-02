"""Ingestão local de currículos recebidos em arquivos para JSON bruto auditável."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import UTC, datetime
from pathlib import Path

from config.settings import CURRICULOS_ENTRADA_DIR, CURRICULOS_PROCESSADOS_DIR
from models.curriculo_extraido import CurriculoExtraido

FORMATOS_SUPORTADOS = {".pdf", ".docx", ".txt", ".md", ".json"}


class ErroExtracaoCurriculo(RuntimeError):
    """Arquivo não pôde ser interpretado localmente."""


def _limpar_texto(texto: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", texto.replace("\r\n", "\n")).strip()


def extrair_texto(path: Path) -> tuple[str, dict[str, str | int | float | bool | None]]:
    """Extrai texto localmente; nenhum conteúdo é transmitido para serviços externos."""
    extensao = path.suffix.lower()
    if extensao in {".txt", ".md", ".json"}:
        return _limpar_texto(path.read_text(encoding="utf-8-sig")), {}
    if extensao == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ErroExtracaoCurriculo("Instale pypdf com: pip install -r requirements.txt") from exc
        reader = PdfReader(str(path))
        texto = "\n".join(pagina.extract_text() or "" for pagina in reader.pages)
        return _limpar_texto(texto), {"paginas": len(reader.pages)}
    if extensao == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise ErroExtracaoCurriculo("Instale python-docx com: pip install -r requirements.txt") from exc
        documento = Document(str(path))
        paragrafos = [paragrafo.text for paragrafo in documento.paragraphs]
        tabelas = [celula.text for tabela in documento.tables for linha in tabela.rows for celula in linha.cells]
        return _limpar_texto("\n".join(paragrafos + tabelas)), {"paragrafos": len(paragrafos), "tabelas": len(documento.tables)}
    raise ErroExtracaoCurriculo(f"Formato não suportado: {extensao or 'sem extensão'}")


def processar_arquivo(path: Path) -> CurriculoExtraido:
    conteudo = path.read_bytes()
    texto, metadados = extrair_texto(path)
    if not texto:
        raise ErroExtracaoCurriculo(
            f"Não foi possível extrair texto de {path.name}. PDFs digitalizados precisam de OCR."
        )
    digest = hashlib.sha256(conteudo).hexdigest()
    return CurriculoExtraido(
        id=f"CUR-{digest[:12].upper()}",
        arquivo_origem=path.name,
        formato=path.suffix.lower().lstrip("."),
        hash_sha256=digest,
        extraido_em=datetime.now(UTC),
        texto=texto,
        metadados=metadados,
    )


def executar_ingestao(
    entrada_dir: Path | None = None,
    output_dir: Path | None = None,
) -> tuple[list[Path], list[str]]:
    """Processa a pasta de entrada e retorna arquivos gerados e avisos."""
    entrada_dir = entrada_dir or CURRICULOS_ENTRADA_DIR
    output_dir = output_dir or CURRICULOS_PROCESSADOS_DIR
    entrada_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    gerados: list[Path] = []
    avisos: list[str] = []
    for path in sorted(entrada_dir.iterdir()):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() not in FORMATOS_SUPORTADOS:
            avisos.append(f"Ignorado {path.name}: formato não suportado.")
            continue
        try:
            curriculo = processar_arquivo(path)
        except ErroExtracaoCurriculo as exc:
            avisos.append(str(exc))
            continue
        destino = output_dir / f"{curriculo.id.lower()}.json"
        destino.write_text(curriculo.model_dump_json(indent=2), encoding="utf-8")
        gerados.append(destino)
    return gerados, avisos


def main() -> None:
    parser = argparse.ArgumentParser(description="Converte currículos em arquivos JSON locais")
    parser.add_argument("--pasta", type=Path, default=CURRICULOS_ENTRADA_DIR, help="Pasta com currículos de entrada")
    parser.add_argument("--saida", type=Path, default=CURRICULOS_PROCESSADOS_DIR, help="Pasta dos JSONs extraídos")
    args = parser.parse_args()
    gerados, avisos = executar_ingestao(args.pasta, args.saida)
    print(f"{len(gerados)} currículo(s) convertido(s) para: {args.saida}")
    for aviso in avisos:
        print(f"Aviso: {aviso}")


if __name__ == "__main__":
    main()
