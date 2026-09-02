"""Gera currículos sintéticos em PDF e DOCX para testar a pasta de entrada."""

from __future__ import annotations

import argparse
import re
import unicodedata
from pathlib import Path

from config.settings import CURRICULOS_ENTRADA_DIR
from generators.curriculos import gerar_curriculos
from models.candidato import Candidato


def _slug(texto: str) -> str:
    normalizado = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "-", normalizado.lower()).strip("-")


def _linhas(curriculo: Candidato) -> list[str]:
    linhas = [
        curriculo.nome,
        f"{curriculo.cargo_desejado} | {curriculo.cidade} - {curriculo.estado}",
        f"E-mail: {curriculo.email} | Telefone: {curriculo.telefone}",
        "",
        "RESUMO PROFISSIONAL",
        curriculo.resumo,
        "",
        "COMPETÊNCIAS",
        ", ".join(curriculo.habilidades),
        "",
        "EXPERIÊNCIA PROFISSIONAL",
    ]
    for experiencia in curriculo.experiencias:
        periodo = f"{experiencia.inicio} - {experiencia.fim or 'Atual'}"
        linhas.extend([f"{experiencia.cargo} | {experiencia.empresa} | {periodo}", experiencia.descricao])
    linhas.extend(["", "FORMAÇÃO"])
    for formacao in curriculo.formacao:
        linhas.append(f"{formacao.curso} - {formacao.instituicao} ({formacao.ano_conclusao or 'em andamento'})")
    if curriculo.certificacoes:
        linhas.extend(["", "CERTIFICAÇÕES", ", ".join(curriculo.certificacoes)])
    if curriculo.idiomas:
        linhas.extend(["", "IDIOMAS", ", ".join(curriculo.idiomas)])
    return linhas


def salvar_pdf(curriculo: Candidato, destino: Path) -> None:
    """Cria um PDF simples, legível e apropriado apenas para dados sintéticos."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("Instale as dependências com: pip install -r requirements.txt") from exc

    estilos = getSampleStyleSheet()
    elementos = []
    for indice, linha in enumerate(_linhas(curriculo)):
        if not linha:
            elementos.append(Spacer(1, 0.18 * cm))
        elif indice == 0:
            elementos.append(Paragraph(f"<b>{linha}</b>", estilos["Title"]))
        elif linha.isupper():
            elementos.append(Paragraph(f"<b>{linha}</b>", estilos["Heading3"]))
        else:
            elementos.append(Paragraph(linha.replace("&", "&amp;"), estilos["BodyText"]))
    SimpleDocTemplate(str(destino), pagesize=A4, leftMargin=1.8 * cm, rightMargin=1.8 * cm).build(elementos)


def salvar_docx(curriculo: Candidato, destino: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Instale as dependências com: pip install -r requirements.txt") from exc

    documento = Document()
    for indice, linha in enumerate(_linhas(curriculo)):
        if not linha:
            documento.add_paragraph()
        elif indice == 0:
            documento.add_heading(linha, level=0)
        elif linha.isupper():
            documento.add_heading(linha, level=1)
        else:
            documento.add_paragraph(linha)
    documento.save(str(destino))


def gerar_documentos(total: int = 5, formato: str = "pdf", output_dir: Path | None = None) -> list[Path]:
    """Cria currículos fictícios independentes dos dados usados no ranking."""
    if formato not in {"pdf", "docx", "ambos"}:
        raise ValueError("Formato inválido. Use pdf, docx ou ambos.")
    output_dir = output_dir or CURRICULOS_ENTRADA_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    curriculos = gerar_curriculos(total=total, seed=2026)
    arquivos: list[Path] = []
    for curriculo in curriculos:
        base = output_dir / f"teste_{_slug(curriculo.nome)}_{curriculo.id.lower()}"
        if formato in {"pdf", "ambos"}:
            path = base.with_suffix(".pdf")
            salvar_pdf(curriculo, path)
            arquivos.append(path)
        if formato in {"docx", "ambos"}:
            path = base.with_suffix(".docx")
            salvar_docx(curriculo, path)
            arquivos.append(path)
    return arquivos


def main() -> None:
    parser = argparse.ArgumentParser(description="Gera currículos sintéticos para testar a ingestão")
    parser.add_argument("--total", type=int, default=5, help="Quantidade de currículos fictícios")
    parser.add_argument("--formato", choices=["pdf", "docx", "ambos"], default="pdf")
    parser.add_argument("--saida", type=Path, default=CURRICULOS_ENTRADA_DIR)
    args = parser.parse_args()
    arquivos = gerar_documentos(args.total, args.formato, args.saida)
    print(f"{len(arquivos)} currículo(s) de teste criado(s) em: {args.saida}")


if __name__ == "__main__":
    main()
